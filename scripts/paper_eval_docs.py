from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer


def _read_json(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        return {}


def _read_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def _read_timeline(path: Path) -> List[Dict[str, Any]]:
    events: List[Dict[str, Any]] = []
    if not path.exists():
        return events
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        raw = line.strip()
        if not raw:
            continue
        try:
            payload = json.loads(raw)
        except Exception:  # noqa: BLE001
            continue
        if isinstance(payload, dict):
            events.append(payload)
    return events


def collect_pipeline_scores(output_dir: Path) -> Dict[str, Any]:
    open_source = _read_json(output_dir / "comparison" / "open_source_initial_score.json")
    logic = _read_json(output_dir / "comparison" / "logic_chain_score.json")
    literature = _read_json(output_dir / "comparison" / "literature_optimization_score.json")
    final_score = _read_json(output_dir / "final_confidence_score.json")
    spec = _read_json(output_dir / "spec.json")
    return {
        "open_source": open_source.get("initial_open_source_score", "n/a"),
        "reproduction": final_score.get("dimension_scores", {}).get("reproduction", spec.get("final_status", "n/a")),
        "logic": logic.get("consistency_score", "n/a"),
        "literature": literature.get("literature_score", "n/a"),
        "overall": final_score.get("confidence_score", "n/a"),
        "final_verdict": final_score.get("final_verdict", "n/a"),
    }


def build_acceptance_markdown(
    *,
    output_dir: Path,
    title: str,
    paper_id: str,
    summary_path: Path,
    timeline_path: Path,
    acceptance_docx_path: Path,
    reproduction_docx_path: Path,
    reproduction_pdf_path: Path,
) -> str:
    summary = _read_json(summary_path)
    timeline = _read_timeline(timeline_path)
    scores = collect_pipeline_scores(output_dir)

    lines = [
        "# Pipeline Final Acceptance Run",
        "",
        f"- Paper ID: {paper_id}",
        f"- Title: {title}",
        f"- Output directory: {output_dir}",
        f"- Last updated: {datetime.now().isoformat(timespec='seconds')}",
        f"- Pipeline status: {summary.get('status', 'running')}",
        "",
        "## Stage Status",
        "",
    ]
    stage_results = summary.get("stage_results", {})
    if stage_results:
        for stage_name, state in stage_results.items():
            lines.append(
                f"- {stage_name}: {state.get('status', 'unknown')} "
                f"(started={state.get('started_at', 'n/a')}, completed={state.get('completed_at', 'n/a')})"
            )
            if state.get("error"):
                lines.append(f"  error: {state['error']}")
    else:
        lines.append("- No stage results yet.")

    lines.extend(
        [
            "",
            "## Scores",
            "",
            f"- Open-source score: {scores['open_source']}",
            f"- Reproduction score/status: {scores['reproduction']}",
            f"- Logic score: {scores['logic']}",
            f"- Literature score: {scores['literature']}",
            f"- Final score: {scores['overall']}",
            f"- Final verdict: {scores['final_verdict']}",
            "",
            "## Artifact Paths",
            "",
            f"- Acceptance DOCX: {acceptance_docx_path}",
            f"- Reproduction report DOCX: {reproduction_docx_path}",
            f"- Reproduction report PDF: {reproduction_pdf_path}",
            "",
            "## Recent Timeline",
            "",
        ]
    )
    if timeline:
        for item in timeline[-20:]:
            lines.append(
                f"- {item.get('timestamp', 'n/a')} | {item.get('stage', 'n/a')} | "
                f"{item.get('action', 'n/a')} | {item.get('status', 'n/a')} | "
                f"{str(item.get('notes', '')).replace(chr(10), ' ')}"
            )
    else:
        lines.append("- No timeline events yet.")
    lines.append("")
    return "\n".join(lines)


def render_markdown_to_docx(markdown_text: str, docx_path: Path, title: str) -> None:
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    document.add_heading(title, level=0)
    in_code = False
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if not stripped:
            document.add_paragraph("")
            continue
        if in_code or stripped.startswith("|"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)
    document.save(str(docx_path))


def render_markdown_to_pdf(markdown_text: str, pdf_path: Path, title: str) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    normal = styles["BodyText"]
    bullet = ParagraphStyle(
        "BulletCompact",
        parent=normal,
        leftIndent=12,
        bulletIndent=0,
        spaceAfter=3,
    )
    code_style = ParagraphStyle(
        "CodeBlock",
        parent=normal,
        fontName="Courier",
        fontSize=8.5,
        leading=10,
        leftIndent=6,
        rightIndent=6,
        spaceAfter=4,
    )

    story: List[Any] = [Paragraph(escape(title), title_style), Spacer(1, 4 * mm)]
    in_code = False
    code_lines: List[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            story.append(Preformatted("\n".join(code_lines), code_style))
            story.append(Spacer(1, 2 * mm))
            code_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code or stripped.startswith("|"):
            code_lines.append(line)
            continue
        flush_code()
        if not stripped:
            story.append(Spacer(1, 2 * mm))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(escape(stripped[2:].strip()), h1))
            continue
        if stripped.startswith("## "):
            story.append(Paragraph(escape(stripped[3:].strip()), h2))
            continue
        if stripped.startswith("- "):
            story.append(Paragraph(escape(stripped[2:].strip()), bullet, bulletText="-"))
            continue
        story.append(Paragraph(escape(stripped), normal))
    flush_code()

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    doc.build(story)
